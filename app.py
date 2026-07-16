import streamlit as st
from openai import OpenAI
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Generator Modul Ajar AI", page_icon="📚", layout="wide")

st.title("Generator Modul Ajar Premium (DOCX Berwarna) 🤖🎨")
st.markdown("Menghasilkan Modul Ajar dengan format tabel **berwarna dan rapi** persis seperti template asli, lengkap hingga lampiran refleksi dan glosarium.")

# --- API KEY ---
try:
    api_key = st.secrets["NVIDIA_API_KEY"]
except KeyError:
    st.error("⚠️ NVIDIA_API_KEY tidak ditemukan di secrets. Pastikan sudah ada di .streamlit/secrets.toml")
    st.stop()

client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",
    api_key=api_key
)

# --- FUNGSI HELPER UNTUK WARNA & TABEL ---
def set_cell_bg(cell, color_hex="D9E2F3"):
    """Fungsi untuk memberi warna background (shading) pada sel tabel Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    tcPr.append(shd)

def style_header_cell(cell, text, bg_color="2F5496", text_color=RGBColor(255, 255, 255)):
    """Memformat sel sebagai Header (Warna gelap, Teks putih tebal)."""
    set_cell_bg(cell, bg_color)
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = text_color

def insert_bullet_points(cell_or_doc, text_data):
    """Mengubah teks dengan baris baru menjadi bullet points. Bisa di dalam sel tabel atau dokumen langsung."""
    lines = str(text_data).split('\n')
    for i, line in enumerate(lines):
        line = line.strip().strip('-').strip('•').strip()
        if line:
            # Cek apakah objek memiliki 'paragraphs' (berarti Cell) atau tidak (berarti Document)
            if hasattr(cell_or_doc, 'paragraphs') and i == 0:
                p = cell_or_doc.paragraphs[0]
            else:
                p = cell_or_doc.add_paragraph()
            p.text = line
            p.style = 'List Bullet'

# --- FUNGSI PEMBUAT DOCX ---
def generate_docx(meta, data):
    doc = Document()
    
    # 1. Judul Utama
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("MODUL AJAR\n")
    r1.bold = True
    r1.font.size = Pt(16)
    r2 = p.add_run("KURIKULUM BERBASIS CINTA – PENDEKATAN DEEP LEARNING\nBerdasarkan Capaian Pembelajaran Terbaru 2025")
    r2.bold = True
    r2.font.size = Pt(12)

    # 2. Tabel Identitas
    doc.add_heading('IDENTITAS MODUL AJAR', level=2)
    t_id = doc.add_table(rows=5, cols=4)
    t_id.style = 'Table Grid'
    
    id_data = [
        ("Mata Pelajaran", meta['mapel'], "Kelas / Fase", meta['kelas']),
        ("Semester", meta['semester'], "Alokasi Waktu", meta['waktu']),
        ("Bab / Topik", meta['topik'], "", ""),
        ("Model Pembelajaran", "PBL (Problem Based Learning)", "Metode Pembelajaran", data.get("metode", "Ceramah, Diskusi")),
        ("Penyusun", meta['guru'], "Sekolah", meta['sekolah'])
    ]
    
    for i, row_data in enumerate(id_data):
        cells = t_id.rows[i].cells
        cells[0].text, cells[1].text = row_data[0], row_data[1]
        set_cell_bg(cells[0], "D9E2F3")
        
        if row_data[2]:
            cells[2].text, cells[3].text = row_data[2], row_data[3]
            set_cell_bg(cells[2], "D9E2F3")
        else:
            cells[1].merge(cells[3])

    doc.add_paragraph()

    # 3. Tabel Identifikasi Peserta Didik
    doc.add_heading('A. IDENTIFIKASI PESERTA DIDIK', level=2)
    t_peserta = doc.add_table(rows=6, cols=2)
    t_peserta.style = 'Table Grid'
    t_peserta.columns[0].width = Pt(150)
    
    peserta_fields = [
        ("Pengetahuan Awal", data.get("pengetahuan_awal", "")),
        ("Minat Belajar", data.get("minat_belajar", "")),
        ("Latar Belakang", data.get("latar_belakang", "")),
        ("Kebutuhan Belajar", data.get("kebutuhan_belajar", "")),
        ("Dimensi Profil Kelulusan", data.get("dimensi_profil", "")),
        ("Topik Panca Cinta", data.get("panca_cinta", ""))
    ]
    
    for i, (k, v) in enumerate(peserta_fields):
        t_peserta.cell(i, 0).text = k
        set_cell_bg(t_peserta.cell(i, 0), "D9E2F3")
        t_peserta.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        insert_bullet_points(t_peserta.cell(i, 1), v)

    doc.add_paragraph()

    # 4. Tabel Desain Pembelajaran
    doc.add_heading('B. DESAIN PEMBELAJARAN', level=2)
    t_desain = doc.add_table(rows=8, cols=2)
    t_desain.style = 'Table Grid'
    
    desain_fields = [
        ("Capaian Pembelajaran (CP)", data.get("cp", "")),
        ("Tujuan Pembelajaran (TP)", data.get("tp", "")),
        ("Lintas Disiplin Ilmu", data.get("lintas_disiplin", "")),
        ("Topik Pembelajaran", data.get("sub_topik", "")),
        ("Praktik Pedagogi", "• Model: PBL\n• Prinsip: Mindful, Meaningful, Joyful"),
        ("Lingkungan Belajar", data.get("lingkungan_belajar", "")),
        ("Kemitraan Pembelajaran", data.get("kemitraan", "")),
        ("Pemanfaatan Digital", data.get("digital", ""))
    ]
    
    for i, (k, v) in enumerate(desain_fields):
        t_desain.cell(i, 0).text = k
        set_cell_bg(t_desain.cell(i, 0), "D9E2F3")
        t_desain.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        insert_bullet_points(t_desain.cell(i, 1), v)
        
    doc.add_paragraph()

    # 5. Tabel Pengalaman Belajar
    doc.add_heading('C. PENGALAMAN BELAJAR', level=2)
    p_pengalaman = doc.add_paragraph()
    p_pengalaman.add_run(f"Materi: {meta['topik']} | Durasi: {meta['waktu']} | Model: PBL").bold = True
    
    t_pengalaman = doc.add_table(rows=4, cols=3)
    t_pengalaman.style = 'Table Grid'
    
    headers = ["FASE KEGIATAN", "AKTIVITAS PEMBELAJARAN", "PRINSIP DL"]
    for i, h in enumerate(headers):
        style_header_cell(t_pengalaman.cell(0, i), h)
        
    fase_data = [
        ("PEMBUKAAN", data.get("kegiatan_pembukaan", ""), "MEANINGFUL\n(Bermakna)"),
        ("MEMAHAMI & MENGAPLIKASIKAN\n(Langkah 1-4 PBL)", data.get("kegiatan_inti", ""), "MINDFUL &\nJOYFUL"),
        ("MEREFLEKSIKAN & PENUTUP", data.get("kegiatan_penutup", ""), "MEANINGFUL")
    ]
    
    for i, (fase, aktivitas, prinsip) in enumerate(fase_data, start=1):
        t_pengalaman.cell(i, 0).text = fase
        set_cell_bg(t_pengalaman.cell(i, 0), "F2F2F2")
        insert_bullet_points(t_pengalaman.cell(i, 1), aktivitas)
        t_pengalaman.cell(i, 2).text = prinsip


    # ==========================================
    # 6. LAMPIRAN I - ASESMEN
    # ==========================================
    doc.add_page_break()
    doc.add_heading('LAMPIRAN I - ASESMEN', level=1)
    
    doc.add_heading('A. Rubrik Penilaian Sikap', level=2)
    insert_bullet_points(doc, data.get("rubrik_sikap", ""))
    
    doc.add_heading('B. Rubrik Penilaian Pengetahuan', level=2)
    insert_bullet_points(doc, data.get("rubrik_pengetahuan", ""))
    
    doc.add_heading('C. Asesmen Sumatif (Soal HOTS)', level=2)
    insert_bullet_points(doc, data.get("soal_hots", ""))

    # ==========================================
    # 7. LAMPIRAN II - MATERI AJAR
    # ==========================================
    doc.add_page_break()
    doc.add_heading('LAMPIRAN II - MATERI AJAR', level=1)
    doc.add_paragraph(data.get("materi_ajar", ""))

    # ==========================================
    # 8. LAMPIRAN III - LKPD
    # ==========================================
    doc.add_page_break()
    doc.add_heading('LAMPIRAN III - LEMBAR KERJA PESERTA DIDIK (LKPD)', level=1)
    doc.add_paragraph(data.get("lkpd", ""))

    # ==========================================
    # 9. LAMPIRAN IV - TINDAK LANJUT DAN REFLEKSI
    # ==========================================
    doc.add_page_break()
    doc.add_heading('LAMPIRAN IV - TINDAK LANJUT DAN REFLEKSI', level=1)
    
    doc.add_heading('A. Program Remedial', level=2)
    insert_bullet_points(doc, data.get("remedial", ""))
    
    doc.add_heading('B. Program Pengayaan', level=2)
    insert_bullet_points(doc, data.get("pengayaan", ""))
    
    doc.add_heading('C. Refleksi Guru', level=2)
    insert_bullet_points(doc, data.get("refleksi_guru", ""))
    
    doc.add_heading('D. Refleksi Peserta Didik', level=2)
    insert_bullet_points(doc, data.get("refleksi_siswa", ""))

    # ==========================================
    # 10. GLOSARIUM & DAFTAR PUSTAKA
    # ==========================================
    doc.add_page_break()
    doc.add_heading('GLOSARIUM', level=1)
    insert_bullet_points(doc, data.get("glosarium", ""))
    
    doc.add_heading('DAFTAR PUSTAKA', level=1)
    doc.add_paragraph(
        "1. Kementerian Pendidikan, Kebudayaan, Riset, dan Teknologi. (2025). Buku Panduan Guru dan Buku Siswa.\n"
        "2. Arends, R.I. (2012). Learning to Teach (9th ed.). New York: McGraw-Hill.\n"
        "3. Barrows, H.S. (1986). A Taxonomy of Problem-Based Learning Methods.\n"
        "4. Referensi relevan lainnya sesuai mata pelajaran."
    )

    # ==========================================
    # 11. TANDA TANGAN
    # ==========================================
    doc.add_paragraph("\n\n")
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.cell(0,0).text = "Mengetahui,\nKepala Sekolah\n\n\n\n\n( ________________________ )"
    sig_table.cell(0,1).text = f"Guru Mata Pelajaran\n\n\n\n\n( {meta['guru']} )"
    
    for row in sig_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


# --- FORM INPUT GURU ---
st.markdown("### 📝 Silakan Isi Data Modul")
with st.form("form_modul"):
    col1, col2 = st.columns(2)
    with col1:
        mata_pelajaran = st.text_input("Mata Pelajaran", placeholder="Contoh: Bahasa Indonesia")
        kelas_fase = st.text_input("Kelas / Fase", placeholder="Contoh: VI / Fase C")
        topik = st.text_input("Bab / Topik Pembelajaran", placeholder="Contoh: Anak-Anak yang Mengubah Dunia")
        semester = st.selectbox("Semester", ["1 (Satu)", "2 (Dua)"])
    with col2:
        nama_guru = st.text_input("Nama Guru Penyusun", placeholder="Contoh: Mamik Muhapatin, S.Pd")
        sekolah = st.text_input("Nama Sekolah", placeholder="Contoh: MI. Miftahussalam")
        alokasi_waktu = st.text_input("Alokasi Waktu", placeholder="Contoh: 24 JP (6 Pertemuan)")

    submit_button = st.form_submit_button("✨ Buat Modul Berwarna Sekarang")

# --- PROSES AI & GENERATE DOCX ---
if submit_button:
    if not mata_pelajaran or not topik:
        st.warning("⚠️ Mohon isi minimal Mata Pelajaran dan Bab/Topik.")
    else:
        with st.spinner("Meracik modul dan mewarnai tabel Word... Mohon tunggu sekitar 15-30 detik."):
            try:
                # Prompt baru yang SAMA SEKALI BUKAN JSON. Super tangguh!
                system_prompt = f"""
                Anda adalah pembuat Modul Ajar SANGAT MENDALAM untuk Topik: "{topik}", Mapel: "{mata_pelajaran}", Fase: "{kelas_fase}".
                
                ATURAN MUTLAK:
                - DILARANG KERAS MENGGUNAKAN FORMAT JSON. Output harus berupa Teks Biasa (Plain Text).
                - Anda WAJIB memisahkan setiap bagian dengan TAG PEMBATAS persis seperti contoh di bawah (contoh: ===METODE===).
                - Gunakan baris baru (Enter) dan tanda strip (-) untuk membuat poin-poin/list.
                
                GUNAKAN TEMPLATE INI (Jawab HANYA dengan format ini dari awal sampai akhir):
                
                ===METODE===
                Ceramah Interaktif, Diskusi Kelompok, Proyek
                ===PENGETAHUAN_AWAL===
                Pengetahuan awal siswa terkait materi...
                ===MINAT_BELAJAR===
                Minat siswa kelas ini pada umumnya...
                ===LATAR_BELAKANG===
                Latar belakang era digital dan keseharian siswa...
                ===KEBUTUHAN_BELAJAR===
                - Visual: ...
                - Audio: ...
                - Kinestetik: ...
                ===DIMENSI_PROFIL===
                - 1. Bernalar Kritis: ...
                - 2. Kreatif: ...
                - 3. Gotong Royong: ...
                ===PANCA_CINTA===
                - 1. Cinta Allah: ...
                - 2. Cinta Sesama: ...
                - 3. Cinta Ilmu: ...
                - 4. Cinta Lingkungan: ...
                - 5. Cinta Tanah Air: ...
                ===CP===
                Tuliskan Capaian Pembelajaran yang relevan...
                ===TP===
                - TP 1 (Pemahaman Dasar): ...
                - TP 2 (Berpikir Kritis): ...
                ===LINTAS_DISIPLIN===
                Kaitan dengan mapel lain (misal IPS/Sains)...
                ===SUB_TOPIK===
                Rincian sub bab yang dibahas...
                ===LINGKUNGAN_BELAJAR===
                Setting kelas dan alat bantu...
                ===KEMITRAAN===
                Guru mapel lain, orang tua...
                ===DIGITAL===
                Canva, Quizizz, Youtube...
                ===KEGIATAN_PEMBUKAAN===
                - Guru mengucapkan salam...
                - Pertanyaan pemantik...
                - Motivasi...
                ===KEGIATAN_INTI===
                - Langkah 1 (Orientasi Masalah): ...
                - Langkah 2 (Organisasi): ...
                - Langkah 3 (Penyelidikan): ...
                ===KEGIATAN_PENUTUP===
                - Kesimpulan...
                - Refleksi...
                - Doa...
                ===RUBRIK_SIKAP===
                - Disiplin: Selalu hadir tepat waktu...
                - Tanggung Jawab: Menyelesaikan tugas...
                - Kerja Sama: Aktif berdiskusi...
                - Toleransi: Menghargai pendapat teman...
                ===RUBRIK_PENGETAHUAN===
                - Mampu mendefinisikan konsep...
                - Mampu menganalisis masalah dengan tepat...
                ===SOAL_HOTS===
                - 1. Soal analisis...
                - 2. Soal evaluasi...
                - 3. Soal kreasi...
                ===MATERI_AJAR===
                Tuliskan ringkasan materi secara mendalam di sini...
                ===LKPD===
                - Tugas 1: ...
                - Tugas 2: ...
                ===REMEDIAL===
                - Bimbingan individu...
                - Pemberian tugas tambahan...
                ===PENGAYAAN===
                - Menjadi tutor sebaya...
                - Mengerjakan proyek lanjutan...
                ===REFLEKSI_GURU===
                - Apa strategi yang paling efektif...
                - Kendala apa yang muncul...
                ===REFLEKSI_SISWA===
                - Bagian mana yang paling disukai...
                - Apa yang masih sulit dipahami...
                ===GLOSARIUM===
                - Istilah 1: Definisi 1...
                - Istilah 2: Definisi 2...
                """

                # Mengirim request ke NVIDIA AI
                completion = client.chat.completions.create(
                    model="nvidia/nemotron-3-ultra-550b-a55b",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": "Tuliskan seluruh modulnya sekarang menggunakan format Tag Pembatas (===KOLOM===). Jangan gunakan JSON."}
                    ],
                    temperature=0.3, # Sedikit dinaikkan agar AI lebih luwes menyusun materi
                    max_tokens=4000 
                )

                hasil_ai = completion.choices[0].message.content
                
                # --- PEMBERSIHAN ANTI-ERROR (TEXT TAGGING PARSER) ---
                # Menggunakan RegEx untuk mengekstrak teks berdasarkan Tag (===NAMA_KOLOM===)
                # Keunggulan metode ini: Mustahil terjadi JSON Error. Jika tulisan terpotong, program tetap jalan!
                data_json = {}
                fields = [
                    "METODE", "PENGETAHUAN_AWAL", "MINAT_BELAJAR", "LATAR_BELAKANG",
                    "KEBUTUHAN_BELAJAR", "DIMENSI_PROFIL", "PANCA_CINTA", "CP", "TP",
                    "LINTAS_DISIPLIN", "SUB_TOPIK", "LINGKUNGAN_BELAJAR", "KEMITRAAN",
                    "DIGITAL", "KEGIATAN_PEMBUKAAN", "KEGIATAN_INTI", "KEGIATAN_PENUTUP",
                    "RUBRIK_SIKAP", "RUBRIK_PENGETAHUAN", "SOAL_HOTS", "MATERI_AJAR",
                    "LKPD", "REMEDIAL", "PENGAYAAN", "REFLEKSI_GURU", "REFLEKSI_SISWA",
                    "GLOSARIUM"
                ]
                
                for field in fields:
                    # Mencari teks mulai dari ===FIELD=== sampai sebelum ===FIELD_SELANJUTNYA=== (atau sampai habis)
                    pattern = fr"==={field}===(.*?)(?====|$)"
                    match = re.search(pattern, hasil_ai, re.DOTALL | re.IGNORECASE)
                    data_json[field.lower()] = match.group(1).strip() if match else ""
                
                # Cek jika AI benar-benar bandel dan tidak memakai tag sama sekali
                if not any(data_json.values()):
                    st.error("⚠️ AI tidak memberikan respons yang sesuai format. Silakan coba lagi.")
                    with st.expander("🔍 Lihat Hasil Teks Asli AI"):
                        st.text(hasil_ai)
                    st.stop()

                # --- LANJUT KE PEMBUATAN DOCX ---
                meta_data = {
                    "mapel": mata_pelajaran,
                    "kelas": kelas_fase,
                    "semester": semester,
                    "waktu": alokasi_waktu,
                    "topik": topik,
                    "guru": nama_guru if nama_guru else "_______________",
                    "sekolah": sekolah if sekolah else "_______________"
                }

                # Generate File Word berwarna
                doc = generate_docx(meta_data, data_json)
                
                # Simpan ke memori untuk tombol download
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.success("✅ Modul Ajar Berwarna (Lengkap dengan seluruh Lampiran) berhasil dibuat!")
                
                # Tombol Download DOCX
                st.download_button(
                    label="📥 Download Modul (Format Word Berwarna)",
                    data=doc_io,
                    file_name=f"Modul_{mata_pelajaran.replace(' ', '_')}_{topik.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"Terjadi kesalahan pada sistem: {e}")
